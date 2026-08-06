from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt


def add_cover_page(document, question):
    """
    Add a professional cover page to the report.
    """

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(18)

    title_run = title.add_run("Radiotherapy Research Assistant")
    title_run.bold = True
    title_run.font.size = Pt(24)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_run = subtitle.add_run(
        "Evidence-Based Research Report"
    )
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(16)

    document.add_paragraph("\n")

    question_heading = document.add_paragraph()
    question_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    question_heading_run = question_heading.add_run(
        "Research Question"
    )
    question_heading_run.bold = True
    question_heading_run.font.size = Pt(13)

    question_paragraph = document.add_paragraph()
    question_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    question_run = question_paragraph.add_run(question)
    question_run.font.size = Pt(12)

    document.add_paragraph("\n")

    generated_paragraph = document.add_paragraph()
    generated_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    generated_run = generated_paragraph.add_run(
        f"Generated: "
        f"{datetime.now().strftime('%d %B %Y, %I:%M %p')}"
    )
    generated_run.font.size = Pt(10)

    version_paragraph = document.add_paragraph()
    version_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    version_run = version_paragraph.add_run(
        "Application Version: 1.0"
    )
    version_run.font.size = Pt(10)

    document.add_page_break()


def add_answer_content(document, answer):
    """
    Add the Claude-generated answer with simple formatting.
    """

    document.add_heading("AI Research Summary", level=1)

    lines = answer.splitlines()

    for line in lines:
        cleaned_line = line.strip()

        if not cleaned_line:
            continue

        # Markdown-style headings
        if cleaned_line.startswith("### "):
            document.add_heading(
                cleaned_line.replace("### ", "", 1),
                level=2
            )

        elif cleaned_line.startswith("## "):
            document.add_heading(
                cleaned_line.replace("## ", "", 1),
                level=2
            )

        elif cleaned_line.startswith("# "):
            document.add_heading(
                cleaned_line.replace("# ", "", 1),
                level=1
            )

        # Bullet points
        elif cleaned_line.startswith("- "):
            paragraph = document.add_paragraph(
                cleaned_line[2:],
                style="List Bullet"
            )

            for run in paragraph.runs:
                run.font.size = Pt(11)

        elif cleaned_line.startswith("• "):
            paragraph = document.add_paragraph(
                cleaned_line[2:],
                style="List Bullet"
            )

            for run in paragraph.runs:
                run.font.size = Pt(11)

        # Numbered sections such as "1. Research Summary"
        elif (
            len(cleaned_line) > 3
            and cleaned_line[0].isdigit()
            and cleaned_line[1:3] == ". "
        ):
            document.add_heading(
                cleaned_line[3:],
                level=2
            )

        else:
            paragraph = document.add_paragraph(cleaned_line)

            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.15

            for run in paragraph.runs:
                run.font.size = Pt(11)


def add_references(document, papers):
    """
    Add retrieved research papers as numbered references.
    """

    document.add_heading("Supporting References", level=1)

    for index, paper in enumerate(papers, start=1):
        reference_table = document.add_table(
            rows=4,
            cols=2
        )

        reference_table.style = "Table Grid"

        labels = [
            "Reference",
            "Title",
            "Authors",
            "Source"
        ]

        values = [
            f"[{index}]",
            paper.get("title", "Unknown"),
            paper.get("authors", "Unknown"),
            (
                f"{paper.get('journal', 'Unknown')} "
                f"({paper.get('year', 'Unknown')})"
            )
        ]

        for row_index in range(4):
            label_cell = reference_table.cell(row_index, 0)
            value_cell = reference_table.cell(row_index, 1)

            label_cell.text = labels[row_index]
            value_cell.text = values[row_index]

            label_cell.vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

            value_cell.vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

            for paragraph in label_cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)

            for paragraph in value_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

        document.add_paragraph()


def add_footer(document):
    """
    Add a footer to every document section.
    """

    for section in document.sections:
        footer = section.footer

        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run(
            "Generated by Radiotherapy Research Assistant v1.0"
        )

        run.font.size = Pt(9)


def configure_document(document):
    """
    Apply document-wide formatting.
    """

    section = document.sections[0]

    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11)

    heading_one = document.styles["Heading 1"]
    heading_one.font.name = "Arial"
    heading_one.font.size = Pt(16)
    heading_one.font.bold = True

    heading_two = document.styles["Heading 2"]
    heading_two.font.name = "Arial"
    heading_two.font.size = Pt(13)
    heading_two.font.bold = True


def export_rag_report(question, answer, papers):
    """
    Export the complete RAG answer and references
    to a professionally formatted Word document.
    """

    output_folder = Path("outputs")
    output_folder.mkdir(exist_ok=True)

    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")

    filename = (
        output_folder
        / f"research_report_{timestamp}.docx"
    )

    document = Document()

    configure_document(document)
    add_cover_page(document, question)
    add_answer_content(document, answer)
    add_references(document, papers)
    add_footer(document)

    document.save(filename)

    return str(filename)