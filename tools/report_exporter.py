from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt


from tools.summarizer import synthesize_papers


def configure_document(document):
    """
    Apply document-wide formatting.
    """

    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11)

    heading_one = document.styles["Heading 1"]
    heading_one.font.name = "Arial"
    heading_one.font.size = Pt(16)
    heading_one.font.bold = True
    heading_one.font.color.rgb = None

    heading_two = document.styles["Heading 2"]
    heading_two.font.name = "Arial"
    heading_two.font.size = Pt(13)
    heading_two.font.bold = True


def add_hyperlink(paragraph, text, url):
    """
    Add a clickable external hyperlink to a paragraph.
    """

    relationship_id = paragraph.part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    run.append(run_properties)

    run_text = OxmlElement("w:t")
    run_text.text = text
    run.append(run_text)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_cover_page(document, query, paper_count):
    """
    Add the report cover page and research-use notice.
    """

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_run = title.add_run("OncoResearch AI")
    title_run.bold = True
    title_run.font.size = Pt(26)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_run = subtitle.add_run(
        "PubMed Oncology Evidence Report"
    )
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(16)

    document.add_paragraph()

    topic_heading = document.add_paragraph()
    topic_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    topic_heading_run = topic_heading.add_run("Research Topic")
    topic_heading_run.bold = True
    topic_heading_run.font.size = Pt(13)

    topic = document.add_paragraph()
    topic.alignment = WD_ALIGN_PARAGRAPH.CENTER

    topic_run = topic.add_run(query)
    topic_run.font.size = Pt(12)

    count = document.add_paragraph()
    count.alignment = WD_ALIGN_PARAGRAPH.CENTER
    count.add_run(f"PubMed records included: {paper_count}")

    generated = document.add_paragraph()
    generated.alignment = WD_ALIGN_PARAGRAPH.CENTER

    generated.add_run(
        "Generated: "
        f"{datetime.now().strftime('%d %B %Y, %I:%M %p')}"
    )

    version = document.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.add_run("Application Version: 1.1.0")

    document.add_paragraph()

    notice = document.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER

    notice_run = notice.add_run(
        "RESEARCH-USE NOTICE\n"
        "This report supports literature research. AI-generated "
        "analyses must be checked against the original publications. "
        "It is not a systematic review or clinical decision-support tool."
    )
    notice_run.bold = True
    notice_run.font.size = Pt(10)

    document.add_page_break()


def add_label_value(document, label, value):
    """
    Add a bold label followed by its value.
    """

    paragraph = document.add_paragraph()

    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True

    value_run = paragraph.add_run(str(value or "Not available"))

    paragraph.paragraph_format.space_after = Pt(6)


def add_source_information(document, paper):
    """
    Add traceable PubMed source information.
    """

    document.add_heading("Source and Verification", level=2)

    add_label_value(
        document,
        "PMID",
        paper.get("pmid", "Not available"),
    )

    add_label_value(
        document,
        "DOI",
        paper.get("doi", "Not available"),
    )

    pubmed_url = paper.get("pubmed_url", "")

    source_paragraph = document.add_paragraph()
    source_label = source_paragraph.add_run("PubMed Source: ")
    source_label.bold = True

    if pubmed_url:
        add_hyperlink(
            source_paragraph,
            pubmed_url,
            pubmed_url,
        )
    else:
        source_paragraph.add_run("Not available")

    add_label_value(
        document,
        "Verification Status",
        paper.get(
            "verification_status",
            "Manual verification required",
        ),
    )


def add_summary_section(document, summary):
    """
    Add the structured AI-assisted analysis.
    """

    document.add_heading("AI-Assisted Analysis", level=2)

    if not isinstance(summary, dict):
        document.add_paragraph(
            "Structured AI analysis was not available."
        )
        return

    add_label_value(
        document,
        "Study Design",
        summary.get("study_design", "Not specified"),
    )

    add_label_value(
        document,
        "Key Findings",
        summary.get("key_findings", "Not specified"),
    )

    add_label_value(
        document,
        "Clinical Significance",
        summary.get(
            "clinical_significance",
            "Not specified",
        ),
    )

    add_label_value(
        document,
        "Limitations",
        summary.get("limitations", "Not specified"),
    )

    keywords = summary.get("keywords", [])

    if isinstance(keywords, list):
        keywords = ", ".join(str(item) for item in keywords)

    add_label_value(
        document,
        "Keywords",
        keywords or "Not specified",
    )


def add_paper(document, paper, index, total_papers):
    """
    Add one PubMed record to the report.
    """

    document.add_heading(
        f"Paper {index}: {paper.get('title', 'Unknown title')}",
        level=1,
    )

    add_label_value(
        document,
        "Authors",
        paper.get("authors", "Unknown"),
    )

    add_label_value(
        document,
        "Journal",
        paper.get("journal", "Unknown"),
    )

    add_label_value(
        document,
        "Year",
        paper.get("year", "Unknown"),
    )

    add_source_information(document, paper)

    document.add_heading("Abstract", level=2)

    abstract = paper.get(
        "abstract",
        "No abstract available in PubMed.",
    )

    abstract_paragraph = document.add_paragraph(str(abstract))
    abstract_paragraph.paragraph_format.space_after = Pt(8)
    abstract_paragraph.paragraph_format.line_spacing = 1.15

    add_summary_section(
        document,
        paper.get("summary", {}),
    )

    if index < total_papers:
        document.add_page_break()


def add_footer(document):
    """
    Add a footer to every document section.
    """

    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run(
            "Generated by OncoResearch AI v1.0 — "
            "Research use only"
        )
        run.font.size = Pt(9)


def export_pubmed_report(query, papers):
    """
    Export PubMed records and AI-assisted analyses to Word.
    """

    output_folder = Path("outputs")
    output_folder.mkdir(exist_ok=True)

    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")

    filename = (
        output_folder
        / f"oncoresearch_pubmed_report_{timestamp}.docx"
    )

    document = Document()

    configure_document(document)
    add_cover_page(document, query, len(papers))

    for index, paper in enumerate(papers, start=1):
        add_paper(
            document=document,
            paper=paper,
            index=index,
            total_papers=len(papers),
        )

    add_footer(document)
    document.save(filename)

    return str(filename)

def export_concise_summary(query, papers):
    """
    Export a concise multi-paper client summary.
    """

    synthesis = synthesize_papers(query, papers)

    output_folder = Path("outputs")
    output_folder.mkdir(exist_ok=True)

    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")

    filename = (
        output_folder
        / f"oncoresearch_concise_summary_{timestamp}.docx"
    )

    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_run = title.add_run(
        "OncoResearch AI\nConcise Evidence Summary"
    )
    title_run.bold = True
    title_run.font.size = Pt(20)

    add_label_value(document, "Research Topic", query)

    add_label_value(
        document,
        "Records Included",
        len(papers),
    )

    add_label_value(
        document,
        "Search Source",
        "PubMed",
    )

    generated_date = datetime.now().strftime("%d %B %Y")

    add_label_value(
        document,
        "Generated",
        generated_date,
    )

    notice = document.add_paragraph()

    notice_run = notice.add_run(
        "Research-use notice: "
        "This is an AI-assisted synthesis of supplied PubMed "
        "records. It is not a systematic review, full-text review "
        "or clinical decision-support report. Findings must be "
        "checked against the cited publications."
    )
    notice_run.bold = True
    notice_run.font.size = Pt(10)

    document.add_heading("Evidence Overview", level=1)

    document.add_paragraph(
        synthesis.get(
            "overview",
            "No overview was generated.",
        )
    )

    document.add_heading("Main Evidence Patterns", level=1)

    evidence_patterns = synthesis.get(
        "evidence_patterns",
        [],
    )

    if evidence_patterns:
        for pattern in evidence_patterns:
            document.add_paragraph(
                str(pattern),
                style="List Bullet",
            )
    else:
        document.add_paragraph(
            "No consistent evidence patterns were identified."
        )

    document.add_heading("Clinical Relevance", level=1)

    document.add_paragraph(
        synthesis.get(
            "clinical_relevance",
            "Not determined.",
        )
    )

    document.add_heading("Evidence Limitations", level=1)

    document.add_paragraph(
        synthesis.get(
            "evidence_limitations",
            "Not determined.",
        )
    )

    document.add_heading("Conclusion", level=1)

    document.add_paragraph(
        synthesis.get(
            "conclusion",
            "No conclusion was generated.",
        )
    )

    document.add_heading("PubMed Records", level=1)

    source_lines = []

    for paper in papers:
        pmid = paper.get("pmid", "Not available")
        source_lines.append(str(pmid))

    document.add_paragraph(
        "PMIDs included: " + ", ".join(source_lines)
    )

    add_footer(document)

    document.save(filename)

    return str(filename)